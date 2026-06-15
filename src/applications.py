"""Generate tailored, editable application materials (resume + cover letter).

Honesty guarantees:
  * The master resume (resume/resume.txt) is read-only and never changed.
  * Resume "tailoring" only REORDERS and EMPHASIZES existing bullets/skills by
    relevance to the job description. It never fabricates employers, dates,
    degrees, certifications, or skills.
  * The cover letter is built from the candidate's real, verifiable background.

Outputs are saved to outputs/applications/<company>_<role>_<date>/.
"""

import re
import sys
import logging
from pathlib import Path
from datetime import datetime

SRC_DIR = Path(__file__).resolve().parent
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

import helpers  # noqa: E402

logger = logging.getLogger("job_assistant.applications")

ROOT_DIR = helpers.ROOT_DIR
OUTPUT_DIR = ROOT_DIR / "outputs" / "applications"
RESUME_FILE = helpers.RESUME_FILE

KNOWN_HEADERS = {
    "PROFESSIONAL EXPERIENCE", "EXPERIENCE", "PROJECTS", "SKILLS",
    "LEADERSHIP & COMMUNITY", "LEADERSHIP", "EDUCATION", "SUMMARY",
}
REORDER_SECTIONS = {"PROFESSIONAL EXPERIENCE", "EXPERIENCE", "PROJECTS"}

# Keywords genuinely supported by the candidate's real background. We only ever
# match/emphasize against these — never injecting unsupported claims.
HONEST_KEYWORDS = [
    "operations", "business operations", "project coordination", "coordination",
    "process improvement", "workflow", "workflow automation", "automation",
    "reporting", "data tracking", "documentation", "technical documentation",
    "python", "api", "apis", "csv", "yaml", "data cleanup", "data", "excel",
    "google sheets", "ai tools", "ai", "communication", "organization",
    "scheduling", "customer service", "customer support", "cross-functional",
    "stakeholder", "business systems", "vs code", "cursor", "problem solving",
    "attention to detail", "onboarding", "analytics", "scheduling",
]

# Real background talking points for the cover letter, mapped to detectable themes.
BACKGROUND_POINTS = [
    ("operations", "coordinating operations and keeping multiple projects on track as an Operations & Administrative Coordinator"),
    ("process improvement", "improving workflows and maintaining organized operational systems"),
    ("automation", "building a Python automation tool that uses APIs, YAML configuration, and CSV pipelines to streamline a repetitive workflow"),
    ("python", "writing Python to automate data collection, cleanup, and reporting"),
    ("data", "working hands-on with CSV/API data, cleanup, and structured YAML configuration"),
    ("ai", "using AI coding tools like Cursor to build and iterate on real software"),
    ("reporting", "tracking performance metrics and compiling reports to support decisions"),
    ("documentation", "producing clear documentation and maintaining professional communication with partners"),
    ("communication", "communicating clearly across teams and with external collaborators"),
]


def _read_master_resume(resume_file=RESUME_FILE):
    try:
        with open(resume_file, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        logger.warning("Master resume not found at %s", resume_file)
        return ""


def _job_keywords(job):
    """Honest keywords that actually appear in the job description/title."""
    text = f"{job.get('title','')} {job.get('description','')}".lower()
    return [kw for kw in HONEST_KEYWORDS if kw in text]


def _relevance(line, job_kws):
    line_lower = line.lower()
    return sum(1 for kw in job_kws if kw in line_lower)


def tailor_resume_text(master_text, job):
    """Return resume text with bullets/skills reordered by job relevance.

    Only reordering happens — no words are added, removed, or invented.
    """
    job_kws = _job_keywords(job)
    if not job_kws or not master_text:
        return master_text

    lines = master_text.splitlines()
    current_section = None
    out = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if stripped.upper() in KNOWN_HEADERS:
            current_section = stripped.upper()
            out.append(line)
            i += 1
            continue

        # Reorder a contiguous block of bullets within experience/project sections.
        if current_section in REORDER_SECTIONS and stripped.startswith(("•", "-", "*")):
            block = []
            while i < n and lines[i].strip().startswith(("•", "-", "*")):
                block.append(lines[i])
                i += 1
            block.sort(key=lambda b: _relevance(b, job_kws), reverse=True)  # stable
            out.extend(block)
            continue

        # Emphasize matching skills by reordering comma-separated skill items.
        if current_section == "SKILLS" and ":" in stripped:
            label, _, items = stripped.partition(":")
            parts = [p.strip() for p in items.split(",") if p.strip()]
            if len(parts) > 1:
                parts.sort(key=lambda p: _relevance(p, job_kws), reverse=True)
                out.append(f"{label}: {', '.join(parts)}")
                i += 1
                continue

        out.append(line)
        i += 1
    return "\n".join(out)


def build_cover_letter_text(job, candidate_name="Ben Goodrum"):
    """Compose a concise, specific cover letter from real background facts."""
    title = job.get("title") or "this role"
    company = job.get("company") or "your team"
    job_kws = _job_keywords(job)

    selected = [phrase for kw, phrase in BACKGROUND_POINTS if kw in job_kws]
    if len(selected) < 3:
        for kw, phrase in BACKGROUND_POINTS:
            if phrase not in selected:
                selected.append(phrase)
            if len(selected) >= 3:
                break
    selected = selected[:3]

    today = datetime.now().strftime("%B %d, %Y")
    greeting = f"Dear {company} Hiring Team,"

    opening = (
        f"I'm writing to apply for the {title} position at {company}. "
        f"As a Business Management graduate with hands-on operations and automation experience, "
        f"I'm confident I can contribute quickly to your team."
    )
    body_intro = "A few things from my background that map directly to this role:"
    bullets = [f"• I have experience {phrase}." for phrase in selected]
    body_close = (
        "Across these experiences I've focused on staying organized, communicating clearly, "
        "and improving the processes around me — the same qualities I'd bring to this position."
    )
    closing = (
        f"I'd welcome the opportunity to discuss how my background fits the {title} role. "
        f"Thank you for your consideration."
    )

    paragraphs = [today, greeting, opening, body_intro]
    paragraphs.extend(bullets)
    paragraphs.append(body_close)
    paragraphs.append(closing)
    paragraphs.append("Sincerely,")
    paragraphs.append(candidate_name)
    return "\n\n".join(paragraphs)


def build_keyword_report(job, master_text):
    """Plain-text report of honest keyword overlap between job and resume."""
    job_kws = _job_keywords(job)
    resume_lower = (master_text or "").lower()
    lines = ["KEYWORD MATCH REPORT", "=" * 22, ""]
    lines.append(f"Job: {job.get('title','')} @ {job.get('company','')}")
    lines.append(f"Honest keywords found in job description: {len(job_kws)}")
    lines.append("")
    lines.append("Keyword                        In Job   In Resume")
    lines.append("-" * 50)
    for kw in HONEST_KEYWORDS:
        in_job = "yes" if kw in job_kws else "-"
        in_resume = "yes" if kw in resume_lower else "-"
        if in_job == "yes" or in_resume == "yes":
            lines.append(f"{kw:<30} {in_job:<8} {in_resume}")
    overlap = [kw for kw in job_kws if kw in resume_lower]
    lines.append("")
    lines.append(f"Aligned keywords (in both job and resume): {len(overlap)}")
    lines.append(", ".join(overlap) if overlap else "(none)")
    return "\n".join(lines)


def _text_to_docx(text, docx_path, title=None):
    """Render plain resume/letter text to a simple, editable .docx."""
    from docx import Document
    from docx.shared import Pt

    document = Document()
    lines = text.splitlines()
    first_real = True
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if first_real:
            heading = document.add_paragraph()
            run = heading.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            first_real = False
            continue
        if stripped.upper() in KNOWN_HEADERS:
            para = document.add_paragraph()
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            continue
        if stripped.startswith(("•", "-", "*")):
            document.add_paragraph(stripped.lstrip("•-* ").strip(), style="List Bullet")
            continue
        document.add_paragraph(stripped)
    document.save(str(docx_path))
    return docx_path


def _application_folder(job):
    company = helpers.sanitize_filename(job.get("company", "Company")) or "Company"
    role = helpers.sanitize_filename(job.get("title", "Role")) or "Role"
    date = datetime.now().strftime("%Y-%m-%d")
    folder = OUTPUT_DIR / f"{company}_{role}_{date}"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def generate_application(job, candidate_name="Ben Goodrum", make_pdf=False):
    """Create tailored resume + cover letter docx (+ previews/report).

    Returns a dict of generated file paths (as strings).
    """
    folder = _application_folder(job)
    master_text = _read_master_resume()
    tailored_resume = tailor_resume_text(master_text, job)
    cover_text = build_cover_letter_text(job, candidate_name=candidate_name)
    report_text = build_keyword_report(job, master_text)

    outputs = {}
    try:
        resume_docx = _text_to_docx(tailored_resume, folder / "tailored_resume.docx")
        outputs["resume_docx"] = str(resume_docx)
    except Exception as exc:
        logger.exception("Failed to write resume docx: %s", exc)

    try:
        cover_docx = _text_to_docx(cover_text, folder / "cover_letter.docx")
        outputs["cover_letter_docx"] = str(cover_docx)
    except Exception as exc:
        logger.exception("Failed to write cover letter docx: %s", exc)

    # Plain-text previews + keyword report (always cheap, always useful).
    try:
        preview_path = folder / "tailored_resume_preview.txt"
        preview_path.write_text(tailored_resume, encoding="utf-8")
        outputs["resume_preview_txt"] = str(preview_path)

        cover_txt_path = folder / "cover_letter_preview.txt"
        cover_txt_path.write_text(cover_text, encoding="utf-8")
        outputs["cover_letter_txt"] = str(cover_txt_path)

        report_path = folder / "keyword_match_report.txt"
        report_path.write_text(report_text, encoding="utf-8")
        outputs["keyword_report_txt"] = str(report_path)
    except Exception as exc:  # pragma: no cover - defensive
        logger.warning("Failed to write preview/report: %s", exc)

    outputs["folder"] = str(folder)
    outputs["cover_letter_text"] = cover_text
    outputs["resume_text"] = tailored_resume
    outputs["keyword_report_text"] = report_text
    return outputs
